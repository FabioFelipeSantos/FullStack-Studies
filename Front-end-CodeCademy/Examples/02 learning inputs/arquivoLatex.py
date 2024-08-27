from docx import Document

# Criar um novo documento Word
doc = Document()

# Título do documento
doc.add_heading('Equações do Documento "Beam.pdf"', 0)

# Lista de equações em LaTeX
equations = [
    ("4.1.1a", "F_V - V - dV + w(x) \\, dx = 0"),
    ("4.1.1b", "\\frac{dV}{dx} = w(x)"),
    ("4.1.1c", "\\frac{dM}{dx} = V"),
    ("4.1.1d", "\\kappa = \\frac{1}{r} = \\frac{M}{EI}"),
    ("4.1.1e", "\\kappa = \\frac{d^2v}{dx^2}"),
    ("4.1.1f", "M = EI \\frac{d^2v}{dx^2}"),
    ("4.1.1g", "\\frac{d^2}{dx^2} \\left(EI \\frac{d^2v}{dx^2}\\right) = w(x)"),
    ("4.1.1h", "EI \\frac{d^4v}{dx^4} = 0"),
    ("4.1.2", "v(x) = a_1 + a_2 x + a_3 x^2 + a_4 x^3"),
    ("4.1.3", "\\begin{aligned} v(0) & = v_1 = a_1 \\\\ \\left.\\frac{dv}{dx}\\right|_{x=0} & = \\theta_1 = a_2 \\\\ v(L) & = v_2 = a_1 + a_2 L + a_3 L^2 + a_4 L^3 \\\\ \\left.\\frac{dv}{dx}\\right|_{x=L} & = \\theta_2 = a_2 + 2a_3 L + 3a_4 L^2 \\end{aligned}"),
    ("4.1.4", "v(x) = N_1 v_1 + N_2 \\theta_1 + N_3 v_2 + N_4 \\theta_2"),
    ("4.1.5", "\\{v\\} = [N] \\{d\\}"),
    ("4.1.6a", "\\{d\\} = \\begin{Bmatrix} v_1 \\\\ \\theta_1 \\\\ v_2 \\\\ \\theta_2 \\end{Bmatrix}"),
    ("4.1.6b", "[N] = [N_1 \\; N_2 \\; N_3 \\; N_4]"),
    ("4.1.7", "\\begin{aligned} N_1 & = 1 - \\frac{3x^2}{L^2} + \\frac{2x^3}{L^3} \\\\ N_2 & = x - \\frac{2x^2}{L} + \\frac{x^3}{L^2} \\\\ N_3 & = \\frac{3x^2}{L^2} - \\frac{2x^3}{L^3} \\\\ N_4 & = - \\frac{x^2}{L} + \\frac{x^3}{L^2} \\end{aligned}"),
    ("4.1.8", "\\epsilon_x = \\frac{du}{dx}"),
    ("4.1.9", "u = -y \\frac{dv}{dx}"),
    ("4.1.10a", "\\epsilon_x = -y \\frac{d^2v}{dx^2}"),
    ("4.1.10b", "\\sigma_x = -y \\frac{d^2v}{dx^2}"),
    ("4.1.11", "\\begin{aligned} m(x) & = EI \\frac{d^2v}{dx^2} \\\\ V & = EI \\frac{d^3v}{dx^3} \\end{aligned}"),
    ("4.1.12", "\\begin{aligned} f_{y1} & = \\frac{EI}{L^3} \\begin{bmatrix} 12 & 6L & -12 & 6L \\end{bmatrix} \\begin{Bmatrix} v_1 \\\\ \\theta_1 \\\\ v_2 \\\\ \\theta_2 \\end{Bmatrix} \\\\ m_1 & = \\frac{EI}{L^3} \\begin{bmatrix} 6L & 4L^2 & -6L & 2L^2 \\end{bmatrix} \\begin{Bmatrix} v_1 \\\\ \\theta_1 \\\\ v_2 \\theta_2 \\end{Bmatrix} \\\\ f_{y2} & = \\frac{EI}{L^3} \\begin{bmatrix} -12 & -6L & 12 & -6L \\end{bmatrix} \\begin{Bmatrix} v_1 \\\\ \\theta_1 \\\\ v_2 \\theta_2 \\end{Bmatrix} \\\\ m_2 & = \\frac{EI}{L^3} \\begin{bmatrix} 6L & 2L^2 & -6L & 4L^2 \\end{bmatrix} \\begin{Bmatrix} v_1 \\theta_1 \\\\ v_2 \\theta_2 \\end{Bmatrix} \\end{aligned}"),
    ("4.1.13", "\\begin{Bmatrix} f_{y1} \\\\ m_1 \\\\ f_{y2} \\\\ m_2 \\end{Bmatrix} = \\frac{EI}{L^3} \\begin{bmatrix} 12 & 6L & -12 & 6L \\\\ 6L & 4L^2 & -6L & 2L^2 \\\\ -12 & -6L & 12 & -6L \\\\ 6L & 2L^2 & -6L & 4L^2 \\end{bmatrix} \\begin{Bmatrix} v_1 \\\\ \\theta_1 \\\\ v_2 \\theta_2 \\end{Bmatrix}"),
    ("4.1.14", "[k] = \\frac{EI}{L^3} \\begin{bmatrix} 12 & 6L & -12 & 6L \\\\ 6L & 4L^2 & -6L & 2L^2 \\\\ -12 & -6L & 12 & -6L \\\\ 6L & 2L^2 & -6L & 4L^2 \\end{bmatrix}"),
    ("4.1.15a", "\\frac{dv}{dx} = \\theta(x) + \\phi(x)"),
]

# Adicionar as equações ao documento
for label, equation in equations:
    doc.add_heading(f'Equação {label}', level=2)
    doc.add_paragraph(equation, style='Normal')

# Salvar o documento
doc_path = "C:/Users/fabio/Downloads/Equacoes_Beam.docx"
doc.save(doc_path)

doc_path
